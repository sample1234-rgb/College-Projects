from tkinter import *
from tkinter import filedialog,font,colorchooser,messagebox
import os,sys,sqlite3
from docx import Document
from fpdf import FPDF
import Designs
'''
class NewCanvas(Canvas):
    def __init__(self,*args,**kargs):
        super(Canvas).__init__(*args,**kargs)

    def create_table(self,x1, y1, x2, y2, r=25, reverse=False, **Options):
        # x1=425, y1=150, x2= 675, y2= 350
        colour_1 = "#ffc466"
        colour_2 = "#ffd485"
        colour_3 = "#ff5656"
        Page.create_rectangle(x1, y1 + r, x2, y2 - r, fill=colour_2, outline="")  # (425,175,675,325) I
        Page.create_rectangle(x1 + r, y2 - r, x2 - r, y2, fill=colour_2, outline="")  # (450,325,650,350) II
        Page.create_oval(x1, y2 - 2 * r, x1 + 2 * r, y2, fill=colour_2, outline="")  # (425,300,475,350) 3
        Page.create_oval(x2 - 2 * r, y2 - 2 * r, x2, y2, fill=colour_2, outline="")  # (625,300,675,350) 2

        Page.create_rectangle(x1 + r, y1, x2 - r, y1 + 2 * r, fill=colour_1, outline="")  # III (450, 200, 650, 150)
        if (reverse):
            Page.create_oval(x1, y1, x1 + 2 * r, y1 + 2 * r, fill=colour_3, outline="")  # 1 (425, 150, 475, 200)
            Page.create_oval(x2 - 2 * r, y1, x2, y1 + 2 * r, fill=colour_1, outline="")  # 4 (625, 150, 675, 200)
        else:
            Page.create_oval(x1, y1, x1 + 2 * r, y1 + 2 * r, fill=colour_1, outline="")  # 1 (425, 150, 475, 200)
            Page.create_oval(x2 - 2 * r, y1, x2, y1 + 2 * r, fill=colour_3, outline="")  # 4 (625, 150, 675, 200)

    def create_rounded_rectangle(self,x1, y1, x2, y2, r=25, color="#ca94ff"):
        Points = [x1, y1 + r, x1 + r, y1 + r, x1 + r, y1, x2 - r, y1, x2 - r, y1 + r, x2, y1 + r, x2, y2 - r, x2 - r,
                  y2 - r, x2 - r, y2, x1 + r, y2, x1 + r, y2 - r, x1, y2 - r]
        colour = color
        Page.create_polygon(Points, fill=colour, outline="")
        Page.create_oval(x1, y1, x1 + 2 * r, y1 + 2 * r, fill=colour, outline="")  # 1 (425, 150, 475, 200)
        Page.create_oval(x2 - 2 * r, y2 - 2 * r, x2, y2, fill=colour, outline="")  # (625,300,675,350) 2
        Page.create_oval(x1, y2 - 2 * r, x1 + 2 * r, y2, fill=colour, outline="")  # (425,300,475,350) 3
        Page.create_oval(x2 - 2 * r, y1, x2, y1 + 2 * r, fill=colour, outline="")  # 4 (625, 150, 675, 200)
'''
class TemplateEditor:
    def __init__(self):
        self.root = Tk()
        self.root.title("Template Editor")
        self.root.minsize(int(0.8 * self.root.winfo_screenwidth()), int(0.9 * self.root.winfo_screenheight()))
        # Panels
        self.panned_window = PanedWindow(bd=5, bg="red", relief=RAISED)
        self.panned_window.pack(fill=BOTH, expand=1)

        self.pan_1 = PanedWindow(bd=0, relief=RAISED)
        self.panned_window.add(self.pan_1, width=125, minsize=50)
        self.frame_1 = Frame(self.pan_1, bg="white")
        self.pan_1.add(self.frame_1)

        self.pan_2 = PanedWindow(bd=0, relief=RAISED)
        self.panned_window.add(self.pan_2, width=900, minsize=900)
        self.frame_2 = Frame(self.pan_2, bg="white")
        self.pan_2.add(self.frame_2)

        self.pan_3 = PanedWindow(bd=0, relief=RAISED)
        self.panned_window.add(self.pan_3, minsize=50)
        self.frame_3 = Frame(self.pan_3, bg="white")
        self.pan_3.add(self.frame_3)

        self.prev_btn = Button(self.frame_2, text="<=", bd=0, command=self.Prev)
        # pan_2.add(prev_btn,minsize=10)
        self.prev_btn.place(relx=0.02, rely=0.95)
        self.Scroll_Bar = Scrollbar(self.frame_2)
        self.Scroll_Bar.pack(side=RIGHT, fill=Y)
        self.Page = Canvas(self.frame_2, bg="white", width=800, height=1000, bd=2)  # ,xscrollcommand=Horizontal_scroll.set,)bd=2)
        self.Page.config(scrollregion=(0, 0, 1000, 1000))
        self.Scroll_Bar.config(command=self.Page.yview)
        # Page.config(width=800,height=1200,yscrollcommand=Scroll_Bar.set,scrollregion=Page.bbox(0,0,1000,1000))
        self.Page.pack()
        # photo_btn = Button(frame_2, text="| O=|' ", bg="white", relief=RIDGE)
        # ds = Designs.Design(Page)
        self.Text_Label_Frame = LabelFrame(self.frame_1, text="Text", bg="lime", fg="white", font=("Ariel", 12, "bold"))
        self.Text_Label_Frame.pack(fill=BOTH, expand=1)  # .place(relx=0.1,rely=0.01)

        bold_btn = Button(self.Text_Label_Frame, text="B", font=("Times New Roman", 10, "bold"), relief=RIDGE)
        bold_btn.grid(row=0, column=0)
        italics_btn = Button(self.Text_Label_Frame, text="I", font=("Times New Roman", 10, "bold"), relief=RIDGE)
        italics_btn.grid(row=0, column=1)
        underline_btn = Button(self.Text_Label_Frame, text="U", font=("Times New Roman", 10), relief=RIDGE)
        underline_btn.grid(row=0, column=2)

        Object_Label_Frame = LabelFrame(self.frame_1, text="Object", bg="yellow", font=("Ariel", 10, "bold"))
        Object_Label_Frame.pack(fill=BOTH, expand=1)  # place(relx=0.1,rely=0.1)

        Text_Box_btn = Button(Object_Label_Frame, text="T", font=("Times New Roman", 10), relief=RIDGE)
        Text_Box_btn.grid(row=0, column=0)
        self.draw_rect_btn = Button(Object_Label_Frame, text="[]", font=("Times New Roman", 10), relief=RIDGE,command=self.Object)
        self.draw_rect_btn.grid(row=0, column=1)
        self.draw_line_btn = Button(Object_Label_Frame, text="/", font=("Times New Roman", 10), relief=RIDGE,command = self.Object)
        self.draw_line_btn.grid(row=0, column=2)
        self.draw_circle_btn = Button(Object_Label_Frame, text="O", font=("Times New Roman", 10), relief=RIDGE,command=self.Object)
        self.draw_circle_btn.grid(row=0, column=3)

        # Panel 3
        File_Lable_Frame = LabelFrame(self.frame_3, text="File", bg="gray", fg="white", font=("Ariel", 12, "bold"))
        File_Lable_Frame.pack(fill=X, expand=1)  # place(relx=0.1,rely=0.01)
        new_btn = Button(File_Lable_Frame, text="New", font=("Times New Roman", 10), relief=RIDGE, command=self.New_File)
        new_btn.pack(fill=X, expand=1)  # padx=5,pady=5)
        open_btn = Button(File_Lable_Frame, text="Open", font=("Times New Roman", 10), relief=RIDGE, command=self.Open_File)
        open_btn.pack(fill=X, expand=1)  # padx=5,pady=5)
        save_btn = Button(File_Lable_Frame, text="Save", font=("Times New Roman", 10), relief=RIDGE, command=self.Save_File)
        save_btn.pack(fill=X, expand=1)  # padx=10,pady=5)
        save_as_btn = Button(File_Lable_Frame, text="Save As", font=("Times New Roman", 10), relief=RIDGE,command=self.Save_As_File)
        save_as_btn.pack(fill=X, expand=1)  # padx=10, pady=5)

        View_Label_Frame = LabelFrame(self.frame_3, text="View", bg="gray", fg="white", font=('Ariel', 12, "bold"))
        View_Label_Frame.pack(fill=X, expand=1)  # place(relx=0.1,rely=0.25)
        Apperance_btn = Button(View_Label_Frame, text="Theme", font=("Times New Roman", 10), relief=RIDGE, command=self.dark)
        Apperance_btn.pack(fill=X, expand=1)

        self.nxt_btn = Button(self.frame_2, text="=>", width=0, bd=0, command=self.Next)
        self.nxt_btn.place(relx=0.95, rely=0.95)
        self.To_Ps("Design_2.ps")
        self.ds = Designs.Design(self.Page)

        self.root.mainloop()
    def Next(self):
        self.nxt_btn.config(command=self.ds.nxt)

        # nxt_btn.config(command=Next)

    def Prev(self):
        self.prev_btn.config(command=self.ds.prev)
        # prev_btn.config(command=Prev)

    def To_Ps(self,file_name):
        # Page.update()
        # if ".ps" not in file_name:
        #     file_name = file_name+".ps"
        self.Page.postscript(file=file_name, colormode="color")

    def To_Pdf(self,file_name):
        self.To_Ps(file_name)
        pdf = FPDF()
        pdf.add_page()
        pdf.output(file_name, "F")

    # New File option
    def New_File(self):
        global opened
        opened = False

        self.Page.delete("all")
        self.root.title("New File")

    def Open_File(self):
        # Delete previous data
        self.Page.delete("all")
        # Geting the file
        file = filedialog.askopenfile(initialdir="C:/Users/ADMIN/Desktop", title="Open",
                                      filetypes=(("Document Files", "*.DOCX"), ("All Files", "*.*")))
        if file:
            global opened
            opened = file
        # updating the status
        file_name = file
        file_name = file_name.replace("C:/", "")
        self.root.title(file_name)
        # Status_bar.config(text=f"{file_name}       ")
        # open the file
        file = open(file, 'r')
        file_data = file.read()
        self.Page.insert(0, file_data)
        # closing the file
        file.close()

    def Save_File(self):
        global opened
        if opened:
            file = open(opened, "w")
            # file.write(Page.get(0, END))
            file.close()
        else:
            self.Save_As_File()

    def Save_As_File(self):
        file = filedialog.asksaveasfile(defaultextension=".*", initialdir="C:", title="Save File", filetypes=(
        ("Text Files", "*.txt"), ("Document Files", "*.DOCX"), ("All Files", "*.*")))
        if file:
            # Update the status
            name = file
            name = name.replace("C:/", "")
            self.root.title(f'{name}')
            # Save the file
            file = open(file, "w")
            # file.write(Page.get(0,END))
            file.close()

    def dark(self):
        back = "#2e2e2e"
        self.frame_1.config(bg=back)
        self.frame_2.config(bg=back)
        self.frame_3.config(bg=back)
        self.panned_window.config(bg="#cccccc")

    def add_rect(self):
        # list1 = [int(x1_entry.get()), int(y1_entry.get()), int(x2_entry.get()), int(y2_entry.get()), fill_entry.get(),outline_entry.get()]
        i = self.Page.create_rectangle(list1[0], list1[1], list1[2], list1[3], fill=list1[4], outline=list1[5])
        btn1 = Button(self.Rect_label, text="text" + str(1))
        btn1.grid(row=4, column=0, columnspan=4)
        self.add_rect_btn.config(text="^", command=lambda i: self.update_rect)

    def update_rect(self,i):
        list1 = [int(x1_entry.get()), int(y1_entry.get()), int(x2_entry.get()), int(y2_entry.get()), fill_entry.get(),
                 outline_entry.get()]
        self.Page.coords(i, list1[0], list1[1], list1[2], list1[3])
        self.Page.itemconfig(i, fill=list1[4], outline=list1[5])
        btn1 = Button(self.Rect_label, text="text" + str(1))
        btn1.grid(row=4, column=0, columnspan=4)
        self.add_rect_btn.config(text=">", command=self.add_rect)

    def Object(self):
        self.Rect_label = LabelFrame(self.frame_1, text="[_]")
        self.Rect_label.pack(fill=BOTH, expand=1)  # place(relx=0.1,rely=0.25)
        # Row 1
        x1_label = Label(self.Rect_label, text="X1:")
        x1_label.grid(row=0, column=0)
        x1_entry = Entry(self.Rect_label, text="X1", width=5)
        x1_entry.grid(row=0, column=1)
        x2_label = Label(self.Rect_label, text="X2:")
        x2_label.grid(row=0, column=2)
        x2_entry = Entry(self.Rect_label, text="X2", width=5)
        x2_entry.grid(row=0, column=3)
        # Row 2
        y1_label = Label(self.Rect_label, text="Y1:")
        y1_label.grid(row=1, column=0)
        y1_entry = Entry(self.Rect_label, text="Y1", width=5)
        y1_entry.grid(row=1, column=1)
        y2_label = Label(self.Rect_label, text="Y2:")
        y2_label.grid(row=1, column=2)
        y2_entry = Entry(self.Rect_label, text="y2", width=5)
        y2_entry.grid(row=1, column=3)
        # Row 3
        fill_label = Label(self.Rect_label, text="Fill:")
        fill_label.grid(row=2, column=0)
        fill_col = colorchooser.askcolor()[1]
        fill_entry = Entry(self.Rect_label, width=5)
        if fill_col =='':
            fill_entry.insert(0, fill_col)
        # fill_entry = Entry(Rect_label, text="White", width=5)
        fill_entry.grid(row=2, column=1)
        outline_label = Label(self.Rect_label, text="Outline:")
        outline_label.grid(row=2, column=2)
        outline_entry = Entry(self.Rect_label, text="Black", width=5)
        outline_entry.grid(row=2, column=3)
        # Row 4
        rect_btn = Button(self.Rect_label, text=" + ")
        rect_btn.grid(row=3, column=2)
        add_rect_btn = Button(self.Rect_label, text=" > ", command=self.add_rect)
        # list1 = [int(x1_entry.get()),int(y1_entry.get()),int(x2_entry.get()),int(y2_entry.get()),fill_entry.get(),outline_entry.get()]
        add_rect_btn.grid(row=3, column=3)
        self.draw_rect_btn.config(command=self.unpack)
        self.draw_line_btn.config(command=self.unpack)
        self.draw_circle_btn.config(command=self.unpack)


    def unpack(self):
        self.Rect_label.pack_forget()
        self.draw_rect_btn.config(command=self.Object)
        self.draw_line_btn.config(command=self.Object)
        self.draw_circle_btn.config(command=self.Object)


'''
# retrieving from database
conn = sqlite3.connect("basic.db")
cursor= conn.cursor()
 x=cursor.execute(''''SELECT * FROM information WHERE first_name = "Gaurav"'''').fetchall()
Z= [y for y in x]
# Variables
fname,lname,email,lindin,no1,no2,add1,add2= Z[0]
x = cursor.execute(''''SELECT * FROM education''''').fetchall()
z = [y for y in x]
edu = z
x = cursor.execute(''''SELECT * FROM work_info'''').fetchall()
skills = cursor.execute(''''SELECT * FROM interests''''').fetchall()
ach = cursor.execute(''''SELECT * FROM achievements'''').fetchall()
conn.close()
global opened,rect,line,oval
opened=False
rect=line=oval=None
# Main Frame
Main_Frame =Frame(root)
Main_Frame.pack(pady=5)
# Scroll Bar
scroll_bar = Scrollbar(Main_Frame)
scroll_bar.pack(side=RIGHT,fill=Y)
# Horizontal scrollbar
Horizontal_scroll=Scrollbar(Main_Frame,orient='horizontal')
Horizontal_scroll.pack(side=BOTTOM,fill=X)
# Canvas
Page=Canvas(Main_Frame,bg="white",width=600,height=600,scrollregion=(0,0,500,400),yscrollcommand=scroll_bar.set,xscrollcommand=Horizontal_scroll.set,bd=2)
# Text box
# Page = Text(Main_Frame,width=80,height=40,undo=True,selectbackground="#fff280",selectforeground="gray",wrap="none",yscrollcommand=scroll_bar.set,xscrollcommand=Horizontal_scroll.set,bd=2)
Page.pack()
# Configuring scroll bar
scroll_bar.config(command=Page.yview)
Horizontal_scroll.config(command=Page.xview)
# Status bar
Status_bar = Label(root,text=" Ready      ",anchor="e")
Status_bar.pack(fill=X,side=BOTTOM,ipady = 2)
# in_frame= Frame(Page,bg="gray",width=300,height=600)
# in_frame.pack(padx=20,pady=20)
'''
'''
def Save_as_Docx_File():
    doc = Document()
    doc.add_heading("Name",1)
    doc.add_heading("Email", 3)
    doc.add_heading("Contact", 3)
    doc.add_heading("Linkedin", 3)
    doc.add_paragraph("Describe yourself")
    doc.add_heading("Education",4)
    doc.add_heading("Achievements",4)
    doc.add_heading("Technical Experience",4)
    doc.add_heading("KEY SKILLS",4)
def colourchoice():
    color = colorchooser.askcolor()[1]
    return color

def Bold():
    B_font=font.Font(Page,Page.cget("font"))
    B_font.config(weight="Bold")

    Page.tag_config("bold",font=B_font)
    current_tags=Page.tag_names("sel.first")
    if "bold" in current_tags:
        Page.tag_remove("bold","sel.first","sel.last")
    else:
        Page.tag_add("bold","sel.first","sel.last")
def Italics():
    I_font = font.Font(Page, Page.cget("font"))
    I_font.config(slant="italic")

    Page.tag_config("italics", font=I_font)
    current_tags = Page.tag_names("sel.first")
    if "italics" in current_tags:
        Page.tag_remove("italics", "sel.first", "sel.last")
    else:
        Page.tag_add("italics", "sel.first", "sel.last")
def Print():
    file = None
    if file:
        # win32ctypes.pywin32.win32api.ShellExecute(0,"print",file,None,".",0)
        pass
'''

'''
# Design 1
def design_1():
    Page.delete("all")
    name= "FIRST LAST NAME"#fname+' '+lname#"FIRST - LAST NAME"
    email="email@email.com"
    lindin="linkedin/username"
    Page.create_text(300,50,text=name,font=('Airel',20,"bold"))
    Page.create_text(120,75,text=email,font=('Helvetica',10))#"Email@email.email"
    no = "xxx-xxx-xxxx"#str(no1)
    # if no2 != 0:
    #     no = str(no1)+",\n"+str(no2)
    Page.create_text(275,75,text="(+91)"+no,font=('Helvetica',10))#"XXX-XXX-XXXX"
    Page.create_text(420,75,text=lindin,font=('Helvetica',10))#"Linkedin@Username"
    Page.create_line(100,100,500,100)#Heading line
    Page.create_text(150,120,text="Describe yourself",font=("Times New Roman",8))
    # Collecting data in list
    # y = 200
    # for i,info in enumerate(edu):
    #     n_info = "> "+info[0]+" "+info[1]+"    "+info[2]#+"     "+str(info[3])
    Page.create_text(300,200,text="Your Education",font=('Airel',8),anchor='w')
    #     y += 20
    Page.create_line(300,170,500,170)#Education qualification list
    Page.create_text(500,160,text="Education",font=('Helvetica',10,"bold"))
    
    Page.create_line(70,170,200,170)#Achievement line
    Page.create_text(70,160,text="Achievements",font=('Airel',10,"bold"))
    # y = 180
    # for i,info in enumerate(ach):
    #     achs = "~ "+info[0]
    Page.create_text(70,180,text="Achievements",font=('Airel',8),anchor="w")
        # y += 20
    Page.create_line(300,330,500,330)#Work experience line
    Page.create_text(500,320,text="Technical Experience",font=('Times New Romans',10,"bold"))
    # if list.length=0 => print("No data")
    # if(len(x) == 0):
    Page.create_text(300,350,text="No Experience till Now",font=('Helvetica',8),anchor="w")
    Page.create_line(50,440,200,440)#interests/hobbbies line
    Page.create_text(50,430,text="KEY SKILLS",font=('Helvetica',10,"bold"))
    # y=460
    # for i,info in enumerate(skills[0]):
    #     tt =''
    #     if (info != '-') or (info != " "):
    #         tt = "> "+info#[0]+" "+info[1]+" "+info[2]+" "+info[3]
    Page.create_text(50,460,text="tt",font=('Airel',8),anchor="w")
        # y += 20
    Page.create_line(100,550,500,550)#About you line
    Page.update()
# Design 2
def design_2():
    Page.delete("all")
    rect_colour="#ffc466"
    Page.create_rectangle(10,0,160,1000,fill=rect_colour,outline="")
    Page.create_rectangle(0,10,170,160,fill=rect_colour,outline="")
    Page.create_rectangle(15,15,155,155,fill="#b0afac")
    name = "FIRST LAST NAME"  # fname+' '+lname#"FIRST - LAST NAME"
    email = "email@email.com"
    lindin = "linkedin/username"
    Page.create_text(80, 175, text=name, font=('Airel', 10, "bold"),justify="center")
    Page.create_text(80, 300, text=email, font=('Helvetica', 10))  # "Email@email.email"
    no = "xxx-xxx-xxxx"  # str(no1)
    # if no2 != 0:
    #     no = str(no1)+",\n"+str(no2)
    Page.create_line(20, 275, 150, 275)  # Heading line(Contacts)
    Page.create_text(80, 325, text="(+91)" + no, font=('Helvetica', 10))  # "XXX-XXX-XXXX"
    Page.create_text(80, 350, text=lindin, font=('Helvetica', 10))  # "Linkedin@Username"
    Page.create_line(20, 185, 150, 185)  # Heading line
    # text= #"Hi, there this is the template of design 2"
    text= "Describe yourself"#, this is the texting text for the wraping effect of canvas create text widget"
    Page.create_text(80, 220, text=text, font=("Times New Roman", 8),width=150)

    Page.create_text(400, 60, text="Education", font=('Helvetica', 12, "bold"))
    Page.create_line(300, 70, 500, 70)  # Education qualification list
    # Collecting data in list
    # y = 200
    # for i,info in enumerate(edu):
    #     n_info = "> "+info[0]+" "+info[1]+"    "+info[2]#+"     "+str(info[3])
    Page.create_text(400, 90, text="Your Education", font=('Airel', 8), anchor='w')
    #     y += 20
    Page.create_text(400, 200, text="Achievements", font=('Airel', 12, "bold"))
    Page.create_line(300, 210, 500, 210)  # Achievement line
    # y = 180
    # for i,info in enumerate(ach):
    #     achs = "~ "+info[0]
    Page.create_text(400, 230, text="--Achievements--", font=('Airel', 8), anchor="w")
    # y += 20
    Page.create_line(300, 380, 500, 380)  # Work experience line
    Page.create_text(400, 370, text="Technical Experience", font=('Times New Romans', 12, "bold"))
    # if list.length=0 => print("No data")
    # if(len(x) == 0):
    Page.create_text(400, 400, text="No Experience till Now", font=('Helvetica', 8), anchor="w")
    Page.create_line(20, 440, 150, 440)  # interests/hobbbies line
    Page.create_text(80, 430, text="KEY SKILLS", font=('Helvetica', 10, "bold"))
    # y=460
    # for i,info in enumerate(skills[0]):
    #     tt =''
    #     if (info != '-') or (info != " "):
    #         tt = "> "+info#[0]+" "+info[1]+" "+info[2]+" "+info[3]
    Page.create_text(50, 460, text="tt", font=('Airel', 8), anchor="w")
    # y += 20
    Page.create_line(160, 550, 800, 550)  # About you line
    Page.update()
# Design 3
def design_3():
    Page.delete("all")
    rect_colour = "#ffc466"
    points_1 =[100,0,0,100,0,120,120,0]
    points_2= [500,600,600,500,600,480,480,600]
    Page.create_polygon(points_1,fill=rect_colour)
    Page.create_polygon(points_2,fill="red")
    # Page.create_rectangle(100, 100, 260, 150, fill=rect_colour, outline="")
    # Page.create_rectangle(0, 10, 170, 160, fill=rect_colour, outline="")
    Page.create_oval(50,50,170,170,outline=rect_colour,width=5)
    Page.create_oval(550,550,450,450,outline="red",width=5)
    name = "FIRST LAST NAME"  # fname+' '+lname#"FIRST - LAST NAME"
    email = "email@email.com"
    lindin = "linkedin/username"
    Page.create_text(250, 40, text=name, font=('Airel', 10, "bold"), justify="center")
    Page.create_text(250, 60, text=email, font=('Helvetica', 10))  # "Email@email.email"
    no = "xxx-xxx-xxxx"  # str(no1)
    # if no2 != 0:
    #     no = str(no1)+",\n"+str(no2)
    Page.create_line(200, 140, 500, 140)  # Heading line
    Page.create_text(250, 80, text="(+91)" + no, font=('Helvetica', 10))  # "XXX-XXX-XXXX"
    Page.create_text(250, 100, text=lindin, font=('Helvetica', 10))  # "Linkedin@Username"
    Page.create_line(325, 45, 325, 125)  # Divider line
    # text= #"Hi, there this is the template of design 2"
    text = "Describe yourself"
    Page.create_text(400, 40, text=text, font=("Times New Roman", 8))

    Page.create_text(500, 190, text="Education", font=('Helvetica', 12, "bold"))
    Page.create_line(250, 200, 500, 200)  # Education qualification list
    # Collecting data in list
    # y = 200
    # for i,info in enumerate(edu):
    #     n_info = "> "+info[0]+" "+info[1]+"    "+info[2]#+"     "+str(info[3])
    Page.create_text(300, 230, text="Your Education", font=('Airel', 8), anchor='w')
    #     y += 20
    Page.create_text(80, 200, text="Achievements", font=('Airel', 12, "bold"))
    Page.create_line(40, 210, 250, 210)  # Achievement line
    # y = 180
    # for i,info in enumerate(ach):
    #     achs = "~ "+info[0]
    Page.create_text(80, 230, text="--Achievements--", font=('Airel', 8), anchor="w")
    # y += 20
    Page.create_line(40, 380, 250, 380)  # Work experience line
    Page.create_text(80, 370, text="Technical Experience", font=('Times New Romans', 12, "bold"))
    # if list.length=0 => print("No data")
    # if(len(x) == 0):
    Page.create_text(40, 400, text="No Experience till Now", font=('Helvetica', 8), anchor="w")
    Page.create_line(250, 400, 500, 400)  # interests/hobbbies line
    Page.create_text(500, 390, text="KEY SKILLS", font=('Helvetica', 10, "bold"))
    # y=460
    # for i,info in enumerate(skills[0]):
    #     tt =''
    #     if (info != '-') or (info != " "):
    #         tt = "> "+info#[0]+" "+info[1]+" "+info[2]+" "+info[3]
    Page.create_text(300, 420, text="tt", font=('Airel', 8), anchor="w")
    # y += 20
    Page.create_line(100, 550, 400, 550,fill="red")  # About you line
    Page.update()
def create_table(x1,y1,x2,y2,r=25,reverse=False,**Options):
    # x1=425, y1=150, x2= 675, y2= 350
    colour_1 = "#ffc466"
    colour_2 = "#ffd485"
    colour_3= "#ff5656"
    Page.create_rectangle(x1,y1+r, x2,y2-r, fill=colour_2, outline="")#(425,175,675,325) I
    Page.create_rectangle(x1+r,y2-r,x2-r,y2, fill=colour_2, outline="")#(450,325,650,350) II
    Page.create_oval(x1,y2-2*r,x1+2*r,y2, fill=colour_2, outline="")#(425,300,475,350) 3
    Page.create_oval(x2-2*r,y2-2*r,x2,y2, fill=colour_2, outline="")#(625,300,675,350) 2

    Page.create_rectangle(x1+r,y1,x2-r,y1+2*r, fill=colour_1, outline="")# III (450, 200, 650, 150)
    if(reverse):
        Page.create_oval(x1,y1,x1+2*r,y1+2*r,fill=colour_3, outline="")# 1 (425, 150, 475, 200)
        Page.create_oval(x2-2*r,y1,x2,y1+2*r, fill=colour_1, outline="")# 4 (625, 150, 675, 200)
    else:
        Page.create_oval(x1, y1, x1 + 2*r, y1 + 2*r, fill=colour_1, outline="")  # 1 (425, 150, 475, 200)
        Page.create_oval(x2 - 2*r, y1, x2, y1 + 2*r, fill=colour_3, outline="")  # 4 (625, 150, 675, 200)
def create_rounded_rectangle(x1,y1,x2,y2,r=25,color="#ca94ff"):
    Points = [x1, y1 + r, x1 + r, y1 + r, x1 + r, y1, x2 - r, y1, x2 - r, y1 + r, x2, y1 + r, x2, y2 - r, x2 - r,
              y2 - r, x2 - r, y2, x1 + r, y2, x1 + r, y2 - r, x1, y2 - r]
    colour=color
    Page.create_polygon(Points, fill=colour, outline="")
    Page.create_oval(x1, y1, x1 + 2 * r, y1 + 2 * r, fill=colour, outline="")  # 1 (425, 150, 475, 200)
    Page.create_oval(x2-2*r,y2-2*r,x2,y2, fill=colour, outline="")#(625,300,675,350) 2
    Page.create_oval(x1,y2-2*r,x1+2*r,y2, fill=colour, outline="")#(425,300,475,350) 3
    Page.create_oval(x2 - 2 * r, y1, x2, y1 + 2 * r, fill=colour, outline="")  # 4 (625, 150, 675, 200)

# Design 4
def design_4():
    Page.delete("all")
    create_rounded_rectangle(20,20,220,800,r=15)
    create_rounded_rectangle(30,30,210,180,r=20,color="#e1c2ff")
    create_rounded_rectangle(30,200,210,350,r=10,color="#cf9eff")
    name = "FIRST LAST NAME"  # fname+' '+lname#"FIRST - LAST NAME"
    email = "email@email.com"
    lindin = "linkedin/username"
    Page.create_text(120, 220, text=name, font=('Airel', 10, "bold"))
    no = "xxx-xxx-xxxx,xxx-xxx-xxxx"  # str(no1)
    # if no2 != 0:
    #     no = str(no1)+",\n"+str(no2)
    create_rounded_rectangle(30,400,210,550,r=15,color="#cf9eff")
    # Page.create_line(200, 140, 500, 140)  # Heading line
    Page.create_text(120, 420, text=email, font=('Helvetica', 10))  # "Email@email.email"
    Page.create_text(120, 460, text="(+91)" + no, font=('Helvetica', 10),width=110)  # "XXX-XXX-XXXX"
    Page.create_text(120, 500, text=lindin, font=('Helvetica', 10))  # "Linkedin@Username"
    Page.create_line(40, 230, 200, 230)  # Divider line
    text= "Decribe yourself Hi, there this is the template of design 2 Describe yourself"
    Page.create_text(120, 260, text=text, font=("Times New Roman", 8),width=150)

    create_table(250,20,770,220)
    Page.create_text(500,40, text="Education", font=('Helvetica', 12, "bold"))
    # Collecting data in list
    # y = 200
    # for i,info in enumerate(edu):
    #     n_info = "> "+info[0]+" "+info[1]+"    "+info[2]#+"     "+str(info[3])
    txt="Hi, there this is the template of design 4 Describe yourself,qwertyuioplkjhgfdsazxcv bnmmmmkiijnhyygvrdzwwwa"
    Page.create_text(400, 80, text=txt, font=('Airel', 8), anchor='w',width=400)
    #     y += 20

    create_table(250,220,770,420,reverse=True)
    Page.create_text(500, 240, text="Achievements", font=('Airel', 12, "bold"))
    # Page.create_line(40, 210, 250, 210)  # Achievement line
    # y = 180
    # for i,info in enumerate(ach):
    #     achs = "~ "+info[0]
    Page.create_text(400, 280, text=txt, font=('Airel', 8), anchor="w",width=400)
    # y += 20

    create_table(250,420,770,620)
    # Page.create_line(40, 380, 250, 380)  # Work experience line
    Page.create_text(500, 440, text="Technical Experience", font=('Times New Romans', 12, "bold"))
    # if list.length=0 => print("No data")
    # if(len(x) == 0):
    Page.create_text(400, 480, text="No Experience till Now", font=('Helvetica', 8), anchor="w",width=400)

    create_table(250,620,770,800,reverse=True)
    # Page.create_line(250, 400, 500, 400)  # interests/hobbbies line
    Page.create_text(500, 640, text="KEY SKILLS", font=('Helvetica', 10, "bold"))
    # y=460
    # for i,info in enumerate(skills[0]):
    #     tt =''
    #     if (info != '-') or (info != " "):
    #         tt = "> "+info#[0]+" "+info[1]+" "+info[2]+" "+info[3]
    Page.create_text(400, 680, text="tt", font=('Airel', 8), anchor="w",width=400)
    # y += 20
    Page.create_line(300, 800, 600, 800, fill="red")  # About you line
    Page.update()
'''
if __name__ == "__main__":
    t = TemplateEditor()

